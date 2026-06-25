import os
import re
import smtplib
import threading
import traceback

from flask import Flask, render_template, request, send_file
from docx import Document
from email.message import EmailMessage

app = Flask(__name__)

# =========================
# CONFIGURACIÓN
# =========================

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

PLANTILLA = os.path.join(BASE_DIR, "plantilla.docx")

CARPETA = os.path.join(BASE_DIR, "documentos_generados")
os.makedirs(CARPETA, exist_ok=True)

# =========================
# CREDENCIALES MAILTRAP
# =========================

EMAIL_USER = os.getenv("EMAIL_USER")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD")

print("EMAIL_USER:", EMAIL_USER)
print("EMAIL_PASSWORD EXISTE:", bool(EMAIL_PASSWORD))

# =========================
# OBTENER VARIABLES
# =========================

def obtener_variables():

    doc = Document(PLANTILLA)

    texto = ""

    for p in doc.paragraphs:
        texto += p.text + "\n"

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                texto += celda.text + "\n"

    variables = re.findall(r"\{\{(.*?)\}\}", texto)

    return sorted(set(variables))

# =========================
# REEMPLAZAR VARIABLES
# =========================

def reemplazar(doc, valores):

    for p in doc.paragraphs:

        for clave, valor in valores.items():

            p.text = p.text.replace(
                "{{" + clave + "}}",
                valor
            )

    for tabla in doc.tables:

        for fila in tabla.rows:

            for celda in tabla.rows[fila._index].cells if False else fila.cells:

                for p in celda.paragraphs:

                    for clave, valor in valores.items():

                        p.text = p.text.replace(
                            "{{" + clave + "}}",
                            valor
                        )

# =========================
# ENVIAR CORREO
# =========================

def enviar_correo(destino, archivo):

    try:

        print("========== ENVIO DE CORREO ==========")
        print("DESTINO:", destino)

        if not EMAIL_USER:
            print("ERROR: EMAIL_USER NO CONFIGURADO")
            return

        if not EMAIL_PASSWORD:
            print("ERROR: EMAIL_PASSWORD NO CONFIGURADO")
            return

        mensaje = EmailMessage()

        mensaje["Subject"] = "Documento generado automáticamente"
        mensaje["From"] = "noreply@mailtrap.io"
        mensaje["To"] = destino

        mensaje.set_content(
            "Adjunto encontrarás el documento generado."
        )

        with open(archivo, "rb") as f:

            mensaje.add_attachment(
                f.read(),
                maintype="application",
                subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename="documento.docx"
            )

        print("Conectando a Mailtrap...")

        with smtplib.SMTP(
            "sandbox.smtp.mailtrap.io",
            2525,
            timeout=60
        ) as smtp:

            smtp.ehlo()
            smtp.starttls()
            smtp.ehlo()

            print("Iniciando sesión Mailtrap...")

            smtp.login(
                EMAIL_USER,
                EMAIL_PASSWORD
            )

            print("Enviando correo...")

            smtp.send_message(mensaje)

        print("CORREO ENVIADO CORRECTAMENTE")

    except Exception:

        print("========== ERROR CORREO ==========")
        print(traceback.format_exc())
        print("==================================")

# =========================
# RUTA PRINCIPAL
# =========================

@app.route("/", methods=["GET", "POST"])
def index():

    try:

        variables = obtener_variables()

        if request.method == "POST":

            print("========== NUEVA SOLICITUD ==========")

            valores = {}

            for v in variables:
                valores[v] = request.form.get(v, "")

            correo = "geovanyasc@gmail.com"

            print("CORREO DESTINO:", correo)

            doc = Document(PLANTILLA)

            reemplazar(doc, valores)

            archivo_salida = os.path.join(
                CARPETA,
                "documento.docx"
            )

            doc.save(archivo_salida)

            print("DOCUMENTO GUARDADO:", archivo_salida)

            if correo:

                print("INICIANDO HILO DE CORREO...")

                threading.Thread(
                    target=enviar_correo,
                    args=(correo, archivo_salida),
                    daemon=True
                ).start()

            return send_file(
                archivo_salida,
                as_attachment=True,
                download_name="documento.docx"
            )

        return render_template(
            "index.html",
            variables=variables
        )

    except Exception as e:

        print("ERROR GENERAL")
        print(traceback.format_exc())

        return (
            "<pre>"
            + str(e)
            + "\n\n"
            + traceback.format_exc()
            + "</pre>",
            500
        )
# =========================
# INICIO
# =========================

if __name__ == "__main__":

    puerto = int(
        os.environ.get(
            "PORT",
            10000
        )
    )

    app.run(
        host="0.0.0.0",
        port=puerto
    )